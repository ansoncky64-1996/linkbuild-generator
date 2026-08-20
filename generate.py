#!/usr/bin/env python3
"""
Digital Zoo Linkbuilding Article Generator
==========================================
讀取 Excel keyword 配對 → AI 生成文章 → 合規檢查 → 產出 .docx / Google Doc

每篇文章喺交付前都要通過 validate_content()(對齊 dz-linkbuild validate.py 嘅規則)。
過唔到就會 repair → 重生成,唔會靜靜地出唔合規嘅稿。

Usage:
    python generate.py --excel data.xlsx --batch 1 --docx out.docx
    python generate.py --excel data.xlsx --batch 1 --single-docx-dir ./out
    python generate.py --excel data.xlsx --batch 1 --creds service_account.json
    python generate.py --excel data.xlsx --batch 1 --dry-run
"""

import os
import sys
import json
import time
import argparse
import logging
import re
import threading
from pathlib import Path

import requests as http_req
from openpyxl import load_workbook
from tqdm import tqdm

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ================================================================
# Configuration
# ================================================================
# ================================================================
# Logging
# ================================================================
# Streamlit Cloud 嘅 stdout 係 pipe,print() 會被 block-buffer(8KB)住,
# 實際上永遠 flush 唔到,所以診斷訊息一定要行 logging(stderr)。
logger = logging.getLogger("dz.linkbuild")
if not logger.handlers:
    _h = logging.StreamHandler(sys.stderr)
    _h.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    logger.addHandler(_h)
    logger.setLevel(logging.INFO)
    logger.propagate = False


def _emit(msg, sink=None, level=logging.INFO):
    """寫落 stderr log,同時收埋落 sink,方便 UI 直接顯示畀用家睇。"""
    logger.log(level, msg)
    if sink is not None:
        sink.append(msg)


OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
MODEL = os.environ.get("LB_MODEL", "~deepseek/deepseek-v4-flash-latest")
DELAY_BETWEEN_CALLS = int(os.environ.get("LB_DELAY", "3"))
# Reasoning model(例如 DeepSeek V4 Flash)嘅推理 token 同 completion token
# 共用 max_tokens。舊值 6000 會俾推理食晒,content 回空 —— 呢個係
# 「API returned empty content」嘅成因。
MAX_TOKENS = int(os.environ.get("LB_MAX_TOKENS", "16000"))
# 除非明確關掉,否則一律叫 provider 唔好做 reasoning。
DISABLE_REASONING = os.environ.get("LB_DISABLE_REASONING", "1") != "0"

# 交付規格(對齊 .claude/agents/dz-linkbuild/scripts/validate.py)
MIN_BODY = 750
MAX_BODY = 1000
TARGET_MIN = 800          # 生成目標留 buffer,避免踩界
TARGET_MAX = 900
H1_LIMIT_CJK = 30
H1_LIMIT_ASCII = 60
N_SECTIONS = 5            # 1 個開篇 + 4 個 H2
LINK_COLOR = "0563C1"     # validator 要求嘅藍色
PLACEHOLDER_URL = "https://example.com"

# 藍色 highlight RGB(Google Docs 0-1 scale)
HIGHLIGHT_COLOR = {"red": 0.60, "green": 0.84, "blue": 0.92}

SCOPES = [
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/drive",
]

CJK_RE = re.compile(r"[㐀-䶿一-鿿豈-﫿]")
PUNCT_RE = re.compile(
    r"[，。、；：？！「」『』（）()《》〈〉—…·．,.;:?!\"'\[\]{}/\\|~@#$%^&*+=<>‧－]"
)
MARKER_RE = re.compile(r"\{\{KW(\d+)\}\}")
DASH_RE = re.compile(r"\s*[—–―]{1,2}\s*")

# 常見簡體字(繁體稿唔應該出現)—— 同 validate.py 一致
SIMPLIFIED = set("们这个国说时来对开会为产业务应该么发过还实点动车马门问题头长间"
                 "东关观规视频网络电脑软风飞马鸟鱼龙进运达远连边过还这样单双证书"
                 "记认识语读写学习医药经济银钱费价买卖务员师专业级别类种样条纸张")
# 大陸用語 → 香港用語。OpenCC 只轉字形唔轉用詞:「视频」會變「視頻」,
# 字形完全正確,簡體檢查放行,但香港客戶要嘅係「影片」。
# 大陸出品嘅 model(Qwen / DeepSeek / 字節 / Moonshot)一定會踩呢啲。
# 只收無歧義嘅詞;有歧義嘅(水平/文件/設置/移動)放落 MAINLAND_WARN 只提示。
MAINLAND_AUTOFIX = {
    "視頻": "影片", "信息": "資訊", "質量": "質素", "服務器": "伺服器",
    "短信": "短訊", "博客": "網誌", "搜索": "搜尋", "屏幕": "螢幕",
    "默認": "預設", "帶寬": "頻寬", "運營商": "電訊商", "雲計算": "雲端運算",
    "算法": "演算法", "數據庫": "資料庫", "接口": "介面", "在線": "網上",
    "移動支付": "流動支付", "移動網絡": "流動網絡", "移動電話": "流動電話",
    "移動裝置": "流動裝置", "移動設備": "流動裝置", "內存": "記憶體",
    "硬盤": "硬碟", "登錄": "登入", "賬號": "帳戶", "賬戶": "帳戶",
    "筆記本電腦": "手提電腦", "小區": "屋苑", "出租車": "的士",
    "公交車": "巴士", "空調": "冷氣", "冰箱": "雪櫃", "概率": "機率",
    "快遞": "速遞", "房價": "樓價", "一次性": "即棄",
}
# 由長到短,避免「移動支付」被「移動網絡」以外嘅短詞搶先
MAINLAND_AUTOFIX_ORDER = sorted(MAINLAND_AUTOFIX, key=len, reverse=True)
# 有歧義,唔敢自動改,只提示
MAINLAND_WARN = {
    "設置": "設定", "水平": "水準", "調研": "調查", "力度": "程度",
    "抓手": "着力點", "落地": "落實", "賦能": "強化", "閉環": "完整流程",
}

# 廣東話口語虛詞(正式 deliverable 唔准出現)
CANTO = ["嘅", "喺", "咗", "嚟", "唔", "冇", "哋", "幾耐", "點算", "邊個", "邊種",
         "鍾意", "而家", "嗰", "咁樣", "嘢", "梗係"]


# ================================================================
# Writing Guidelines (injected into AI prompt)
# ================================================================
GUIDELINES_EN = f"""
## DZ Linkbuilding Compliance Ruleset

### Length
- {MIN_BODY}-{MAX_BODY} words of body content (headings excluded). Aim for {TARGET_MIN}-{TARGET_MAX}.
- Each of the five sections must carry roughly {TARGET_MIN // N_SECTIONS}-{TARGET_MAX // N_SECTIONS} words.

### Content Principles
- The thematic angle must be broader than either keyword, so the article never competes with the client's own target pages.
- Write as online media content: natural, editorial, never promotional.

### Heading Rules
- One H1, under {H1_LIMIT_ASCII} characters, no punctuation of any kind, and it must not contain any keyword.
- Exactly four H2s. No H3s. No keyword may appear in any heading.
- Structure: opening paragraph (no H2), then four H2 sections.

### Keyword Placement
- Each keyword appears exactly once in the body, written only as its marker.
- Keywords must not appear in any heading, in the opening paragraph, or in the closing paragraph.
- The two markers must be separated by at least one complete keyword-free section.
- The keyword wording, including partial matches, must not appear anywhere else. Paraphrase instead.

### Style
- Bullets may supplement narrative paragraphs but never replace them.
- No em-dashes or en-dashes. Use commas, full stops, or restructure the sentence.
"""

GUIDELINES_ZH = f"""
## DZ 外連文章合規規則

### 字數
- 正文 {MIN_BODY}–{MAX_BODY} 個漢字(不含標題),目標 {TARGET_MIN}–{TARGET_MAX} 字。
- 五個 section 每個約 {TARGET_MIN // N_SECTIONS}–{TARGET_MAX // N_SECTIONS} 字,唔可以有段落明顯偏短。

### 內容原則
- 主題角度必須闊過任何一個關鍵字,避免同客戶自己嘅目標頁面互搶排名。
- 以網絡媒體編輯內容嘅方式寫,自然、有觀點,唔可以似廣告。

### 標題規則
- 一個 H1,{H1_LIMIT_CJK} 個全形字以內,完全唔可以有標點符號,亦唔可以包含關鍵字。
- 剛好四個 H2,冇 H3。所有標題都唔可以出現關鍵字。
- 結構:開篇段落(冇 H2)+ 四個 H2 section。

### 關鍵字擺位
- 每個關鍵字喺正文剛好出現一次,而且只可以用標記形式出現。
- 關鍵字唔可以出現喺任何標題、開篇段落或結尾段落。
- 兩個標記之間要隔住至少一整個完全冇關鍵字嘅 section。
- 關鍵字字面(包括部分匹配)唔可以喺其他地方出現,要用同義講法改寫。

### 風格
- 可以用列點輔助,但唔可以用列點取代論述段落。
- 唔准用破折號(—、–),改用逗號、句號或者改寫句子。
"""


# ================================================================
# Small helpers
# ================================================================
def _as_int(value):
    """把 Excel cell 轉做 int。文字格式嘅 '1' / '1.0' / ' 2 ' 都收。轉唔到回 None。"""
    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value) if value.is_integer() else None
    try:
        return int(float(str(value).strip()))
    except (TypeError, ValueError):
        return None


def _as_text(value):
    """把 Excel cell 轉做乾淨字串,數字/日期都唔會炸。"""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    return str(value).strip()


def _kw_pattern(kw):
    """一律唔分大細階。

    中英混合嘅關鍵字(「mpls 虛擬專用網絡」)之前行 case-sensitive,
    model 寫「MPLS 虛擬專用網絡」就 match 唔到,marker 貼唔落去。
    大細階折疊對中文字冇影響,所以全部 IGNORECASE 係安全嘅。
    """
    return re.compile(re.escape(kw), re.IGNORECASE)


def _kw_count(haystack, kw):
    if not kw or not haystack:
        return 0
    return len(_kw_pattern(kw).findall(haystack))


def _fragment_hit(keyword, haystack):
    """關鍵字冇完整 match,但有 ≥2/3 長度嘅連續片段出現。"""
    if not keyword or not haystack:
        return None
    if keyword.isascii():
        keyword, haystack = keyword.lower(), haystack.lower()
    n = max(2, int(len(keyword) * 2 / 3))
    for i in range(len(keyword) - n + 1):
        frag = keyword[i:i + n]
        if frag in haystack:
            return frag
    return None


def _keyword_slots(article):
    """回 [(marker, keyword, url)],由長到短排 —— 長嘅關鍵字要優先認領,
    否則 '紅酒' 會搶咗 '紅酒櫃' 入面嗰三個字。"""
    slots = []
    for idx in (1, 2):
        kw = article.get(f"keyword{idx}", "").strip()
        if kw:
            slots.append((
                "{{KW%d}}" % idx,
                kw,
                article.get(f"url{idx}", "").strip() or PLACEHOLDER_URL,
            ))
    slots.sort(key=lambda s: len(s[1]), reverse=True)
    return slots


def _mask_markers(text):
    """把 {{KWn}} 換做同長度嘅哨兵,令 keyword 掃描唔會撞到 marker 內容。"""
    return MARKER_RE.sub(lambda m: "\x00" * len(m.group(0)), text)


def _render_body(body, article):
    """把 marker 換返真正嘅關鍵字文字(用嚟做字數統計 / 交畀 validator)。"""
    out = body
    for idx in (1, 2):
        kw = article.get(f"keyword{idx}", "").strip()
        out = out.replace("{{KW%d}}" % idx, kw)
    return MARKER_RE.sub("", out)


# ================================================================
# Excel Parser
# ================================================================
def parse_excel(filepath, batch_number, sheet_name=None, report=None):
    """讀 Excel,回傳指定 batch 嘅文章 list。

    容錯重點:
      * 序號欄可以係數字或者文字格式('1' / 1 / 1.0 都收)
      * Batch 標籤大細階同空白都唔敏感('Batch 1' / 'batch1' / 'BATCH  1')
      * 任何讀唔到嘅行都會記錄落 report,唔會靜靜地消失
    """
    warnings = report if report is not None else []
    wb = load_workbook(filepath, data_only=True)
    if sheet_name is None:
        sheet_name = wb.sheetnames[0]
    if sheet_name not in wb.sheetnames:
        wb.close()
        raise KeyError(
            f"Excel 冇 sheet「{sheet_name}」,現有:{', '.join(wb.sheetnames)}"
        )
    ws = wb[sheet_name]

    def batch_key(label):
        return re.sub(r"\s+", "", str(label)).lower()

    target_key = batch_key(f"Batch {batch_number}")

    articles = []
    current_batch = None

    row = 2  # skip header
    while row <= ws.max_row:
        batch_val = ws[f"A{row}"].value
        if batch_val not in (None, ""):
            current_batch = batch_key(batch_val)

        if current_batch != target_key:
            row += 1
            continue

        art_num = _as_int(ws[f"B{row}"].value)
        kw_num = _as_int(ws[f"C{row}"].value)
        kw_cell = _as_text(ws[f"D{row}"].value)

        # 有關鍵字但序號讀唔到 —— 一定要嘈,唔可以當佢冇出現過
        if kw_num is None and kw_cell and kw_cell != "--":
            warnings.append(
                f"第 {row} 行有關鍵字「{kw_cell}」,但 C 欄序號讀唔到"
                f"(值 = {ws[f'C{row}'].value!r}),已跳過"
            )
            row += 1
            continue

        if art_num is not None and kw_num == 1:
            kw1 = kw_cell
            url1 = _as_text(ws[f"E{row}"].value)
            category = _as_text(ws[f"K{row}"].value) or "General"

            kw2 = ""
            url2 = ""
            if row + 1 <= ws.max_row and _as_int(ws[f"C{row+1}"].value) == 2:
                kw2_val = _as_text(ws[f"D{row+1}"].value) or "--"
                url2_val = _as_text(ws[f"E{row+1}"].value) or "--"
                if kw2_val != "--":
                    kw2 = kw2_val
                if url2_val != "--":
                    url2 = url2_val

            art_warnings = []
            if not kw1:
                warnings.append(f"文章 #{art_num}:keyword 1 係空白,已跳過")
                row += 1
                continue
            if not url1:
                art_warnings.append(f"keyword 1「{kw1}」冇 target URL,暫用 placeholder")
                url1 = PLACEHOLDER_URL
            if kw2 and not url2:
                art_warnings.append(f"keyword 2「{kw2}」冇 target URL,暫用 placeholder")
                url2 = PLACEHOLDER_URL
            if not kw2:
                art_warnings.append("只有 1 個 keyword(DZ 標準係每篇 2 個),請確認 batch brief")

            articles.append({
                "number": art_num,
                "keyword1": kw1,
                "url1": url1,
                "keyword2": kw2,
                "url2": url2,
                "category": category,
                "warnings": art_warnings,
            })

        row += 1

    wb.close()

    seen = {}
    for a in articles:
        seen.setdefault(a["number"], []).append(a)
    for num, group in seen.items():
        if len(group) > 1:
            warnings.append(f"文章 #{num} 喺 Batch {batch_number} 出現咗 {len(group)} 次,請檢查 Excel")

    return articles


# ================================================================
# Language Detection
# ================================================================
def detect_language(text):
    """由關鍵字判斷文章語言。

    規則:關鍵字含任何中文字 → 中文稿;純英文關鍵字先至出英文稿。
    """
    return "zh-HK" if CJK_RE.search(text or "") else "en"


# ================================================================
# Prompt Building
# ================================================================
def _kw_section(article, lang):
    kw1, url1 = article["keyword1"], article["url1"]
    kw2, url2 = article["keyword2"], article["url2"]

    if lang == "zh-HK":
        if kw2:
            return (
                f"目標關鍵字:\n"
                f"- 關鍵字 1:「{kw1}」(目標連結:{url1})\n"
                f"- 關鍵字 2:「{kw2}」(目標連結:{url2})\n\n"
                f"文章要同時觸及兩個關鍵字嘅主題範疇。"
                f"用 {{{{KW1}}}} 標記關鍵字 1 嘅位置,用 {{{{KW2}}}} 標記關鍵字 2 嘅位置。"
            )
        return (
            f"目標關鍵字:\n- 關鍵字 1:「{kw1}」(目標連結:{url1})\n\n"
            f"用 {{{{KW1}}}} 標記關鍵字 1 嘅位置。呢篇只有一個關鍵字,"
            f"絕對唔可以出現 {{{{KW2}}}}。"
        )

    if kw2:
        return (
            f"Target keywords:\n"
            f'- Keyword 1: "{kw1}" (target link: {url1})\n'
            f'- Keyword 2: "{kw2}" (target link: {url2})\n\n'
            f"The article must cover the subject areas of both keywords. "
            f"Write {{{{KW1}}}} where keyword 1 belongs and {{{{KW2}}}} where keyword 2 belongs."
        )
    return (
        f'Target keywords:\n- Keyword 1: "{kw1}" (target link: {url1})\n\n'
        f"Write {{{{KW1}}}} where keyword 1 belongs. This article has only one keyword, "
        f"so {{{{KW2}}}} must never appear."
    )


def build_prompt(article):
    """組 AI prompt。中文稿全中文 prompt,英文稿全英文 prompt ——
    兩邊唔可以撈埋,否則 model 會跟住 prompt 語言出錯文章語言。"""
    lang = detect_language(article["keyword1"] + article["keyword2"])
    category = article["category"]
    kw_block = _kw_section(article, lang)

    if lang == "zh-HK":
        return f"""你是一位專業嘅 SEO 內容寫手。請根據以下指引撰寫一篇反向連結文章。

{GUIDELINES_ZH}

---

### 本篇任務

{kw_block}

網站類別 / 口吻:{category}

### 語言要求
- 必須使用繁體中文(香港標準字形),嚴禁出現任何簡體字
- 用繁體中文書面語,語體對標香港主流網絡媒體(例如《香港01》、《經濟日報》副刊)
- 嚴禁廣東話口語同粵語語助詞(「嘅」「咗」「啲」「點解」「攞」「揀」「搞掂」「嚟」等)
- 用「的」不用「嘅」,用「了」不用「咗」,用「一些」不用「啲」,用「為什麼」不用「點解」
- 嚴禁破折號(—、–、――),改用逗號、句號或者改寫句子
- 用詞正式但自然,段落之間要有邏輯銜接

### 輸出格式(嚴格遵守)

只輸出一個 JSON object,冇任何其他文字,冇 code fence。
剛好 {N_SECTIONS} 個 sections:1 個開篇(h2 為 null)+ 4 個 H2 section。

{{
  "h1": "標題,{H1_LIMIT_CJK} 全形字以內,完全冇標點,唔含關鍵字",
  "sections": [
    {{"h2": null, "body": "開篇段落,唔含關鍵字"}},
    {{"h2": "第一個 H2 標題", "body": "正文段落,可以喺呢度嵌入 {{{{KW1}}}}"}},
    {{"h2": "第二個 H2 標題", "body": "緩衝段落,唔含任何關鍵字"}},
    {{"h2": "第三個 H2 標題", "body": "正文段落,可以喺呢度嵌入 {{{{KW2}}}}"}},
    {{"h2": "第四個 H2 標題", "body": "結尾段落,唔含關鍵字"}}
  ]
}}

### 硬性規則(違反任何一條都會被退回重寫)
1. 每個 marker 各只出現一次,而且一定要用 {{{{KW1}}}} / {{{{KW2}}}} 呢個標記形式
2. 關鍵字字面(包括部分匹配)唔可以喺標記以外任何位置出現,要用同義講法改寫
3. 兩個標記之間要隔住至少一整個冇關鍵字嘅 section
4. 標記唔可以放喺第一個 section(開篇)或最後一個 section(結尾)
5. H1 同所有 H2 都唔可以包含關鍵字或者佢嘅任何片段
6. 一定要剛好 {N_SECTIONS} 個 sections,唔可以多亦唔可以少
7. 文章主題唔可以直接等於關鍵字,要搵一個更上層嘅主題角度
8. 正文總字數(不含標題)必須喺 {TARGET_MIN}–{TARGET_MAX} 個漢字之間。
   每個 section 只可以寫 {TARGET_MIN // N_SECTIONS}–{TARGET_MAX // N_SECTIONS} 個漢字,
   五段加埋**絕對唔可以超過 {MAX_BODY} 字**。寧願寫少啲,唔好寫多。
   寫之前先數清楚,唔好一路寫一路加內容
9. H1 唔可以有任何標點符號,包括頓號、冒號、問號、引號
"""

    return f"""You are a professional SEO content writer. Write one backlink article following the ruleset below.

{GUIDELINES_EN}

---

### This assignment

{kw_block}

Site category / tone: {category}

### Language requirements
- Write in fluent, professional English suitable for quality online media.
- Formal but accessible. Smooth logical transitions between paragraphs.
- Do NOT use em-dashes or en-dashes. Use commas, full stops, or restructure.
- Every single word of output must be English. Do not use any Chinese characters anywhere.

### Output format (strict)

Output one JSON object only. No surrounding prose, no code fence.
Exactly {N_SECTIONS} sections: one opening (h2 = null) plus four H2 sections.

{{
  "h1": "Title, under {H1_LIMIT_ASCII} characters, no punctuation at all, no keyword",
  "sections": [
    {{"h2": null, "body": "Opening paragraph, no keyword"}},
    {{"h2": "First H2 title", "body": "Body paragraph, may embed {{{{KW1}}}} here"}},
    {{"h2": "Second H2 title", "body": "Buffer paragraph, no keyword at all"}},
    {{"h2": "Third H2 title", "body": "Body paragraph, may embed {{{{KW2}}}} here"}},
    {{"h2": "Fourth H2 title", "body": "Closing paragraph, no keyword"}}
  ]
}}

### Hard rules (breaking any one sends the draft back)
1. Each marker appears exactly once, always in the {{{{KW1}}}} / {{{{KW2}}}} marker form.
2. The literal keyword wording, including partial matches, must not appear anywhere outside its marker. Paraphrase instead.
3. The two markers must be separated by at least one complete keyword-free section.
4. No marker in the first section (opening) or the last section (closing).
5. Neither the H1 nor any H2 may contain a keyword or any fragment of one.
6. Exactly {N_SECTIONS} sections, no more and no fewer.
7. The article topic must not simply restate a keyword. Choose a broader angle.
8. Total body word count excluding headings MUST land between {TARGET_MIN} and {TARGET_MAX} words.
   Each section may run only {TARGET_MIN // N_SECTIONS}-{TARGET_MAX // N_SECTIONS} words, and the five
   together must NEVER exceed {MAX_BODY} words. Err on the short side, never the long side.
9. The H1 must contain no punctuation of any kind, including colons, question marks and quotes.
"""


# ================================================================
# Normalisation
# ================================================================
_OPENCC = None
_OPENCC_LOCK = threading.Lock()

# OpenCC 嘅 s2hk 除咗轉簡體之外,仲會把本身完全正確嘅繁體字改成
# 香港教育局字形變體(客戶→客户、說明→説明、稅務→税務、溫度→温度…)。
# 呢啲字形喺香港主流網絡媒體係唔會咁寫嘅,而且 validate.py 嘅簡體字表
# 亦捉唔到,所以要喺轉換之後還原返標準寫法。
HK_GLYPH_RESTORE = {
    "户": "戶", "説": "說", "悦": "悅", "閲": "閱", "鋭": "銳", "税": "稅",
    "脱": "脫", "温": "溫", "兑": "兌", "羣": "群", "敍": "敘", "愠": "慍",
    "蜕": "蛻", "藴": "蘊", "醖": "醞", "卧": "臥", "葱": "蔥", "緼": "縕",
}
_HK_GLYPH_TABLE = str.maketrans(HK_GLYPH_RESTORE)


def _to_traditional(text):
    """把簡體字轉做香港繁體,同時保住標準繁體字形。"""
    if not text:
        return text
    global _OPENCC
    if _OPENCC is None:
        with _OPENCC_LOCK:
            if _OPENCC is None:
                try:
                    from opencc import OpenCC
                    _OPENCC = OpenCC("s2hk")
                except Exception:
                    _OPENCC = False
    if not _OPENCC:
        return text.translate(_HK_GLYPH_TABLE)
    try:
        return _OPENCC.convert(text).translate(_HK_GLYPH_TABLE)
    except Exception:
        return text.translate(_HK_GLYPH_TABLE)


def _remove_dashes(text, lang):
    """破折號換返正常標點,按文章語言決定用全形逗號定半形。"""
    if not text:
        return text
    replacement = "，" if lang == "zh-HK" else ", "
    text = DASH_RE.sub(replacement, text)
    text = re.sub(r"\s+--\s+", replacement, text)
    return text


def _mainland_autofix(text, protect=()):
    """大陸用語換返香港用語,但唔可以郁到關鍵字本身。

    客戶關鍵字可能就係「視頻會議系統」,改咗個 keyword 就永遠 match 唔返,
    所以先把關鍵字出現位置遮住,改完再還原。
    """
    if not text:
        return text

    spans = []

    def stash(m):
        spans.append(m.group(0))
        return f"\x01{len(spans) - 1}\x02"

    masked = text
    for kw in sorted((k for k in protect if k), key=len, reverse=True):
        masked = _kw_pattern(kw).sub(stash, masked)

    for src in MAINLAND_AUTOFIX_ORDER:
        masked = masked.replace(src, MAINLAND_AUTOFIX[src])

    return re.sub(r"\x01(\d+)\x02", lambda m: spans[int(m.group(1))], masked)


def _normalize_output(result, lang, article=None):
    """繁體轉換 + 破折號清理。

    一定要喺 keyword 處理之前跑:model 好常出簡體,
    如果先做 keyword 比對再轉繁,簡體寫嘅關鍵字會走漏,
    轉繁之後就變成一個冇 hyperlink 嘅重複關鍵字。
    """
    protect = ()
    if article:
        protect = (article.get("keyword1", ""), article.get("keyword2", ""))

    def clean(s):
        if not s:
            return s
        s = _remove_dashes(s, lang)
        if lang == "zh-HK":
            s = _to_traditional(s)
            s = _mainland_autofix(s, protect)
        return s

    result["h1"] = clean(result.get("h1", ""))
    for section in result.get("sections", []):
        if section.get("h2"):
            section["h2"] = clean(section["h2"])
        if section.get("body"):
            section["body"] = clean(section["body"])
    return result


# ================================================================
# Keyword marker placement (non-destructive)
# ================================================================
def _place_markers(result, article):
    """確保每個關鍵字剛好有一個 marker。

    同舊版最大分別:呢度唔會再用 re.sub 剷走文字。剷字會整爛句子
    (「購置紅酒櫃」→「購置櫃」、"many wine cellars"→"many s"),
    亦會令長關鍵字被短關鍵字食咗。剩低嘅 bare occurrence 交由
    validate_content() 報告,再由 _repair_article() 叫 model 改寫。
    """
    sections = result["sections"]
    n = len(sections)
    slots = _keyword_slots(article)

    # 1. 每個 marker 全篇只可以保留一個(留第一個)
    for marker, _kw, _url in slots:
        seen = False
        for section in sections:
            body = section.get("body", "")
            if marker not in body:
                continue
            if not seen:
                first = body.index(marker)
                head = body[:first + len(marker)]
                tail = body[first + len(marker):].replace(marker, "")
                section["body"] = head + tail
                seen = True
            else:
                section["body"] = body.replace(marker, "")

    # 2. 缺 marker 嘅,由 bare occurrence 升格。長關鍵字行先。
    for marker, kw, _url in slots:
        if any(marker in s.get("body", "") for s in sections):
            continue

        # 中段優先(開篇同結尾唔准放關鍵字)
        for idx in list(range(1, max(1, n - 1))):
            body = sections[idx].get("body", "")
            masked = _mask_markers(body)
            m = _kw_pattern(kw).search(masked)
            if not m:
                continue
            sections[idx]["body"] = body[:m.start()] + marker + body[m.end():]
            break
        # 搵唔到就唔會硬塞。validate_content() 會報 marker 缺失,
        # 然後由 repair / 重生成處理。

    return result


def _strip_unmapped_markers(result, article):
    """清走冇對應關鍵字嘅 marker,例如單關鍵字文章入面嘅 {{KW2}}。
    唔清嘅話個 marker 會原封不動印落 .docx 交去 client 手上。"""
    valid = {marker for marker, _kw, _url in _keyword_slots(article)}
    dropped = 0

    def scrub(text):
        nonlocal dropped

        def repl(m):
            nonlocal dropped
            if m.group(0) in valid:
                return m.group(0)
            dropped += 1
            return ""

        text = MARKER_RE.sub(repl, text)
        return re.sub(r" {2,}", " ", text)

    result["h1"] = scrub(result.get("h1", ""))
    for section in result["sections"]:
        if section.get("h2"):
            section["h2"] = scrub(section["h2"])
        if section.get("body"):
            section["body"] = scrub(section["body"])
    return result, dropped


# ================================================================
# Counting
# ================================================================
def _count_words(result, lang, article=None):
    """數正文字數。Marker 會先還原做真正嘅關鍵字文字,
    咁計出嚟先同 validator 睇到嘅成品一致。"""
    parts = []
    for section in result.get("sections", []):
        body = section.get("body", "")
        parts.append(_render_body(body, article) if article else MARKER_RE.sub("", body))
    text = " ".join(parts)

    if lang == "en":
        return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'’\-]*", text))
    return len(CJK_RE.findall(text))


# ================================================================
# Compliance validation (mirrors dz-linkbuild validate.py)
# ================================================================
def validate_content(result, article, lang):
    """交付前合規檢查。回 (fails, warns) 兩個 list。

    規則對齊 .claude/agents/dz-linkbuild/scripts/validate.py,
    但係喺 JSON 階段跑,唔使等 build 完 docx 先發現。
    """
    fails, warns = [], []
    sections = result.get("sections", [])

    # ---------- 1. 結構 ----------
    if len(sections) != N_SECTIONS:
        fails.append(f"Section 數量:應該 {N_SECTIONS} 個,實際 {len(sections)} 個")
        return fails, warns
    if sections[0].get("h2") is not None:
        fails.append("第一個 section 必須係開篇(h2 = null)")
    missing_h2 = [i for i in range(1, N_SECTIONS) if not sections[i].get("h2")]
    if missing_h2:
        fails.append(f"Section {missing_h2} 缺 H2 標題")
    empty = [i for i, s in enumerate(sections) if not (s.get("body") or "").strip()]
    if empty:
        fails.append(f"Section {empty} 冇正文")
    if fails:
        return fails, warns

    h1 = result.get("h1", "").strip()
    heads_text = " ".join([h1] + [s.get("h2") or "" for s in sections])
    bodies = [s.get("body", "") for s in sections]

    # ---------- 2. H1 ----------
    if not h1:
        fails.append("冇 H1")
    else:
        if CJK_RE.search(h1):
            length, limit, unit = len(re.sub(r"\s", "", h1)), H1_LIMIT_CJK, "全形字元"
        else:
            length, limit, unit = len(h1), H1_LIMIT_ASCII, "ASCII 字元"
        if length >= limit:
            fails.append(f"H1 長度:{length} {unit},上限 under {limit}")
        found_punct = sorted(set(PUNCT_RE.findall(h1)))
        if found_punct:
            fails.append(f"H1 唔准有標點:{' '.join(found_punct)}")

    # ---------- 3. 語言 ----------
    body_text_raw = "".join(bodies)
    full_raw = body_text_raw + heads_text
    if lang == "zh-HK":
        simp = sorted(set(full_raw) & SIMPLIFIED)
        if simp:
            fails.append("偵測到簡體字:" + " ".join(simp))
        canto = [w for w in CANTO if w in full_raw]
        if canto:
            fails.append("偵測到廣東話口語虛詞:" + " ".join(canto))
        # 關鍵字本身可能就係大陸用語(例如客戶指定「視頻會議系統」),
        # 嗰啲唔算漏網,掃之前要遮住。
        vocab_scan = full_raw
        for _m, _kw, _u in _keyword_slots(article):
            vocab_scan = _kw_pattern(_kw).sub("", vocab_scan)
        mainland = [f"{w}→{MAINLAND_WARN[w]}" for w in MAINLAND_WARN if w in vocab_scan]
        if mainland:
            warns.append("疑似大陸用語(要人手判斷):" + " ".join(mainland))
        leaked = [f"{w}→{MAINLAND_AUTOFIX[w]}" for w in MAINLAND_AUTOFIX_ORDER
                  if w in vocab_scan]
        if leaked:
            warns.append("大陸用語自動替換有漏網:" + " ".join(leaked))
        variants = sorted(set(full_raw) & set(HK_GLYPH_RESTORE))
        if variants:
            warns.append(
                "偵測到非主流字形變體:"
                + " ".join(f"{v}→{HK_GLYPH_RESTORE[v]}" for v in variants)
            )
    else:
        stray = CJK_RE.findall(full_raw)
        if stray:
            fails.append(
                f"英文稿入面有 {len(stray)} 個中文字:" + " ".join(sorted(set(stray))[:12])
            )
    if DASH_RE.search(full_raw):
        fails.append("偵測到破折號(—/–),規格禁止")

    # ---------- 4. Marker ----------
    slots = _keyword_slots(article)
    marker_section = {}
    for marker, kw, url in slots:
        idxs = [i for i, b in enumerate(bodies) if marker in b]
        total = sum(b.count(marker) for b in bodies)
        if total == 0:
            fails.append(f"「{kw}」:正文搵唔到 marker {marker}")
            continue
        if total > 1:
            fails.append(f"「{kw}」:marker 出現 {total} 次,規格要求剛好 1 次")
        marker_section[marker] = idxs[0]
        if idxs[0] == 0:
            fails.append(f"「{kw}」:出現喺開篇段落,規格禁止")
        elif idxs[0] == len(bodies) - 1:
            fails.append(f"「{kw}」:出現喺結尾段落,規格禁止")
        if url == PLACEHOLDER_URL:
            warns.append(f"「{kw}」仲用緊 placeholder URL,交稿前要換返真實 target URL")

    stray_markers = {
        m.group(0) for b in bodies + [h1, heads_text] for m in MARKER_RE.finditer(b)
    } - {s[0] for s in slots}
    if stray_markers:
        fails.append("有未對應關鍵字嘅 marker:" + " ".join(sorted(stray_markers)))

    # ---------- 5. Buffer ----------
    if len(marker_section) == 2:
        a, b = sorted(marker_section.values())
        if a == b:
            fails.append("Keyword buffer:兩個關鍵字喺同一段")
        elif b - a < 2:
            fails.append("Keyword buffer:兩個關鍵字喺相鄰段落,中間要有一整段 buffer")

    # ---------- 6. 關鍵字字面外洩 ----------
    masked_bodies = _mask_markers(body_text_raw)
    for _marker, kw, _url in slots:          # 已由長到短排,長嘅先食
        n = _kw_count(masked_bodies, kw)
        if n:
            fails.append(
                f"「{kw}」:除咗 marker 之外,正文仲有 {n} 次字面出現,要改寫做同義表達"
            )
            masked_bodies = _kw_pattern(kw).sub("\x00", masked_bodies)

    masked_heads = _mask_markers(heads_text)
    for _marker, kw, _url in slots:
        if _kw_count(masked_heads, kw):
            fails.append(f"「{kw}」:出現喺 H1 或 H2,規格禁止")
        else:
            frag = _fragment_hit(kw, masked_heads)
            if frag:
                warns.append(f"「{kw}」:標題有近似片段「{frag}」,建議換講法")

    # ---------- 7. 字數 ----------
    n = _count_words(result, lang, article)
    unit = "英文字" if lang == "en" else "中文字"
    if n < MIN_BODY:
        fails.append(f"正文字數:{n} {unit},唔夠 {MIN_BODY},要補 {MIN_BODY - n}")
    elif n > MAX_BODY:
        fails.append(f"正文字數:{n} {unit},超出 {MAX_BODY},要刪 {n - MAX_BODY}")

    return fails, warns


# ================================================================
# OpenRouter plumbing
# ================================================================
ARTICLE_SCHEMA = {
    "name": "linkbuild_article",
    "strict": True,
    "schema": {
        "type": "object",
        "additionalProperties": False,
        "required": ["h1", "sections"],
        "properties": {
            "h1": {"type": "string"},
            "sections": {
                "type": "array",
                "minItems": N_SECTIONS,
                "maxItems": N_SECTIONS,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["h2", "body"],
                    "properties": {
                        "h2": {"type": ["string", "null"]},
                        "body": {"type": "string"},
                    },
                },
            },
        },
    },
}


def _extract_json(raw):
    """由 model 回覆抽出 JSON object,容忍 code fence 同前後多餘文字。"""
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    # strict=False:容忍字串入面嘅原始換行(Invalid control character)
    try:
        return json.loads(raw, strict=False)
    except json.JSONDecodeError:
        start, end = raw.find("{"), raw.rfind("}")
        if start == -1 or end <= start:
            raise
        return json.loads(raw[start:end + 1], strict=False)


def _chat(api_key, model, prompt, max_tokens=None, temperature=0.7,
          schema=None, log=None):
    """打 OpenRouter。

    兩個關鍵設定:
      * reasoning 預設關掉 —— reasoning model 嘅推理 token 同 completion
        共用 max_tokens,唔關就成日回空 content。
      * 支援 structured outputs 嘅 model 會用 json_schema,直接杜絕
        「唔係合法 JSON」同「section 數量唔啱」。
    provider 唔收邊個參數就會自動除返佢再試,唔會成篇失敗。
    """
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens or MAX_TOKENS,
    }
    if DISABLE_REASONING:
        payload["reasoning"] = {"enabled": False}
    if schema:
        payload["response_format"] = {"type": "json_schema", "json_schema": schema}

    optional = ["reasoning", "response_format"]

    while True:
        resp = http_req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=180,
        )

        if resp.status_code >= 400:
            try:
                detail = resp.json().get("error", {}).get("message", "")
            except Exception:
                detail = resp.text[:300]
            # provider 唔支援某個參數就除返佢再試一次
            dropped = next(
                (k for k in optional
                 if k in payload and k.replace("_", " ") in detail.lower().replace("_", " ")),
                None,
            )
            if dropped:
                _emit(f"provider 唔收 {dropped},除咗佢再試:{detail[:120]}", log,
                      logging.WARNING)
                payload.pop(dropped)
                optional.remove(dropped)
                continue
            raise http_req.exceptions.HTTPError(
                f"OpenRouter {resp.status_code}: {detail} (model = {model!r})"
            )

        data = resp.json()
        choice = (data.get("choices") or [{}])[0]
        message = choice.get("message") or {}
        content = message.get("content")

        if content:
            return content

        # 空 content:一定要講清楚點解,唔好再盲
        usage = data.get("usage") or {}
        detail_bits = [
            f"finish_reason={choice.get('finish_reason')}",
            f"native={choice.get('native_finish_reason')}",
            f"completion_tokens={usage.get('completion_tokens')}",
            f"max_tokens={payload['max_tokens']}",
        ]
        reasoning = message.get("reasoning") or ""
        if reasoning:
            detail_bits.append(f"reasoning_chars={len(reasoning)}")
        if message.get("refusal"):
            detail_bits.append(f"refusal={message['refusal'][:100]}")

        # 推理食晒 budget:再加大一次 max_tokens 重試
        if (reasoning or choice.get("finish_reason") == "length") \
                and payload["max_tokens"] < 48000:
            payload["max_tokens"] = min(payload["max_tokens"] * 3, 48000)
            _emit("content 回空(" + ", ".join(detail_bits)
                  + f"),max_tokens 加到 {payload['max_tokens']} 重試", log,
                  logging.WARNING)
            continue

        raise ValueError("API 回空 content(" + ", ".join(detail_bits) + ")")


def _assert_shape(result):
    assert "h1" in result, "Missing h1"
    assert "sections" in result, "Missing sections"
    assert len(result["sections"]) == N_SECTIONS, \
        f"Need exactly {N_SECTIONS} sections (got {len(result['sections'])})"
    assert result["sections"][0].get("h2") is None, "First section must have h2=null"
    for idx in range(1, N_SECTIONS):
        assert result["sections"][idx].get("h2"), f"Section {idx} missing H2 title"


def _repair_article(result, article, lang, problems, api_key, model, log=None):
    """把合規問題列表交返畀 model 修。改唔到就回原本嗰份(由外層決定重生成)。

    Prompt 語言跟返文章語言 —— 舊版無論中英文都用中文 prompt 叫佢改,
    英文稿改完會變中文。
    """
    current = json.dumps(result, ensure_ascii=False, indent=None)
    issue_list = "\n".join(f"{i+1}. {p}" for i, p in enumerate(problems))

    if lang == "zh-HK":
        prompt = f"""以下 JSON 文章未通過合規檢查。請修正所有問題後,輸出完整嘅修訂版 JSON。

### 必須修正嘅問題
{issue_list}

### 修改要求
- 保持剛好 {N_SECTIONS} 個 sections 嘅結構(1 個開篇 h2 = null + 4 個 H2)
- {{{{KW1}}}} / {{{{KW2}}}} 標記各保留剛好一個,唔可以刪走
- 如果問題係「字面出現」,要把嗰啲字眼改寫做同義講法,唔好純粹刪走令句子唔通
- 如果問題係字數,要真係增補或者刪減論述內容,目標 {TARGET_MIN}–{TARGET_MAX} 字
- 全篇必須繁體中文書面語,冇簡體字、冇廣東話口語、冇破折號
- 只輸出 JSON,唔要 code fence,唔要任何解釋文字

原文 JSON:
{current}"""
    else:
        prompt = f"""The JSON article below failed compliance review. Fix every issue and return the complete revised JSON.

### Issues that must be fixed
{issue_list}

### Requirements
- Keep exactly {N_SECTIONS} sections (one opening with h2 = null, then four H2 sections).
- Keep exactly one {{{{KW1}}}} and one {{{{KW2}}}} marker. Never delete a marker.
- Where the issue is a literal keyword occurrence, paraphrase that wording. Do not simply delete words and leave a broken sentence.
- Where the issue is word count, genuinely expand or tighten the argument. Target {TARGET_MIN}-{TARGET_MAX} words.
- The entire output must be English, with no em-dashes or en-dashes.
- Output JSON only. No code fence, no commentary.

Original JSON:
{current}"""

    try:
        raw = _chat(api_key, model, prompt, temperature=0.5,
                    schema=ARTICLE_SCHEMA, log=log)
        repaired = _extract_json(raw)
        _assert_shape(repaired)
        return repaired
    except Exception as e:
        _emit(f"修正失敗:{type(e).__name__}: {e}", log, logging.WARNING)
        return result


def generate_article_content(article, api_key, model, max_retries=3,
                             max_repairs=2, log=None, return_best_effort=True):
    """生成一篇合規文章。過唔到 validate_content() 就唔會回稿。

    log:傳一個 list 入嚟就會收齊所有診斷訊息,UI 可以直接顯示,
    唔使叫用家去掘 server log。
    """
    if log is None:
        log = []
    prompt = build_prompt(article)
    lang = detect_language(article.get("keyword1", "") + article.get("keyword2", ""))
    num = article.get("number", "?")
    last_fails = []
    best = None          # 過唔到 gate 但最接近嘅一份,唔好白白掉咗
    best_score = None

    for attempt in range(max_retries):
        try:
            raw = _chat(api_key, model, prompt, temperature=0.7,
                        schema=ARTICLE_SCHEMA, log=log)
            result = _extract_json(raw)
            _assert_shape(result)

            for repair_round in range(max_repairs + 1):
                # 次序好重要:先轉繁 → 先擺 marker → 再清 stray marker → 最後驗
                result = _normalize_output(result, lang, article)
                result = _place_markers(result, article)
                result, dropped = _strip_unmapped_markers(result, article)
                if dropped:
                    _emit(f"#{num} 清走咗 {dropped} 個無效 marker", log, logging.WARNING)

                fails, warns = validate_content(result, article, lang)
                if not fails:
                    result["_warnings"] = warns
                    result["_word_count"] = _count_words(result, lang, article)
                    result["_lang"] = lang
                    result["_compliant"] = True
                    result["_log"] = list(log)
                    return result

                last_fails = fails
                if best_score is None or len(fails) < best_score:
                    best, best_score = json.loads(json.dumps(result)), len(fails)
                    best["_fails"] = list(fails)
                if repair_round == max_repairs:
                    break
                _emit(f"#{num} 合規未過({len(fails)} 項),第 {repair_round + 1} 次修正", log, logging.WARNING)
                for f in fails:
                    _emit(f"    • {f}", log, logging.WARNING)
                result = _repair_article(
                    result, article, lang, fails, api_key, model, log=log
                )

            raise AssertionError(
                f"修正 {max_repairs} 次後仍有 {len(last_fails)} 項未過"
            )

        except (json.JSONDecodeError, AssertionError, KeyError, ValueError) as e:
            _emit(f"#{num} 第 {attempt+1} 次嘗試失敗:{type(e).__name__}: {e}", log, logging.WARNING)
            if attempt < max_retries - 1:
                time.sleep(5)
            continue
        except http_req.exceptions.RequestException as e:
            _emit(f"#{num} 第 {attempt+1} 次嘗試 API 出錯:{type(e).__name__}: {e}", log, logging.ERROR)
            if attempt < max_retries - 1:
                time.sleep(10)
            continue

    _emit(f"#{num} 試咗 {max_retries} 次都唔得", log, logging.ERROR)
    if last_fails:
        _emit("最後一次未過嘅項目:", log, logging.ERROR)
        for f in last_fails:
            _emit(f"    • {f}", log, logging.ERROR)

    if best is not None and return_best_effort:
        # 過唔到 gate,但有一份接近嘅草稿。明確標示做未合規交返出去,
        # 好過成篇冇 —— 由用家決定人手執定重跑。
        best["_compliant"] = False
        best["_warnings"] = []
        best["_word_count"] = _count_words(best, lang, article)
        best["_lang"] = lang
        best["_log"] = list(log)
        _emit(f"#{num} 交返一份未合規草稿({best_score} 項未過),需要人手處理",
              log, logging.ERROR)
        return best
    return None


# ================================================================
# Document Builder (Google Docs)
# ================================================================
class DocBuilder:
    """Builds document text and tracks formatting ranges."""

    def __init__(self):
        self.text = ""
        self.bold_ranges = []       # (start, end)
        self.keyword_ranges = []    # (start, end, url)

    @property
    def pos(self):
        return len(self.text)

    def add(self, content, bold=False):
        start = self.pos
        self.text += content
        if bold:
            self.bold_ranges.append((start, self.pos))
        return start

    def add_line(self, content, bold=False):
        start = self.add(content, bold=bold)
        self.text += "\n"
        return start

    def add_blank(self):
        self.text += "\n"

    def add_keyword(self, keyword_text, url):
        start = self.pos
        self.text += keyword_text
        self.keyword_ranges.append((start, self.pos, url))

    def build_article(self, article, content):
        num = article["number"]
        kw1, kw2 = article["keyword1"], article["keyword2"]
        url1, url2 = article["url1"], article["url2"]

        self.add_line(f"#{num}", bold=True)
        self.add_blank()

        self.add_line("Keyword：", bold=True)
        self.add_blank()
        self.add_line(f"● {kw1}")
        if kw2:
            self.add_line(f"● {kw2}")
        self.add_blank()

        # 標題格式對齊 validate.py 嘅 [Hn: 標題文字]
        self.add_line(f"[H1: {content['h1']}]", bold=True)
        self.add_blank()

        for section in content["sections"]:
            h2 = section.get("h2")
            if h2:
                self.add_line(f"[H2: {h2}]", bold=True)
                self.add_blank()
            self._add_body_with_keywords(
                section.get("body", ""), kw1, url1, kw2, url2
            )
            self.add_blank()

        self.add_blank()

    def _add_body_with_keywords(self, body, kw1, url1, kw2, url2):
        parts = re.split(r"(\{\{KW[12]\}\})", body)
        for part in parts:
            if part == "{{KW1}}" and kw1:
                self.add_keyword(kw1, url1 or PLACEHOLDER_URL)
            elif part == "{{KW2}}" and kw2:
                self.add_keyword(kw2, url2 or PLACEHOLDER_URL)
            elif MARKER_RE.fullmatch(part or ""):
                continue          # 未對應嘅 marker 一律唔印
            else:
                self.text += part
        self.text += "\n"

    def get_gdocs_requests(self):
        requests = [{
            "insertText": {"location": {"index": 1}, "text": self.text}
        }]
        for start, end in self.bold_ranges:
            requests.append({
                "updateTextStyle": {
                    "range": {"startIndex": start + 1, "endIndex": end + 1},
                    "textStyle": {"bold": True},
                    "fields": "bold",
                }
            })
        for start, end, url in self.keyword_ranges:
            requests.append({
                "updateTextStyle": {
                    "range": {"startIndex": start + 1, "endIndex": end + 1},
                    "textStyle": {
                        "backgroundColor": {"color": {"rgbColor": HIGHLIGHT_COLOR}},
                        "foregroundColor": {
                            "color": {"rgbColor": {
                                "red": 0x05 / 255, "green": 0x63 / 255, "blue": 0xC1 / 255,
                            }}
                        },
                        "underline": True,
                        "link": {"url": url},
                    },
                    "fields": "backgroundColor,foregroundColor,underline,link",
                }
            })
        return requests


# ================================================================
# DOCX File Builder
# ================================================================
def _set_run_font(run, size=11, bold=False, color=None):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")


def _add_hyperlink_run(paragraph, text, url):
    """插入真 hyperlink,藍色 0563C1 + 單底線(validator 會逐項檢查)。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    url = (url or "").strip() or PLACEHOLDER_URL   # 空 URL 會令 Word 報檔案損毀

    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    rpr.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_COLOR)
    rpr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    rpr.append(size)

    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "cyan")
    rpr.append(hl)

    run_el.append(rpr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)

    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


def _add_body_paragraph(doc, body, article):
    p = doc.add_paragraph()
    parts = re.split(r"(\{\{KW[12]\}\})", body)
    for part in parts:
        if part == "{{KW1}}" and article["keyword1"]:
            _add_hyperlink_run(p, article["keyword1"], article["url1"])
        elif part == "{{KW2}}" and article["keyword2"]:
            _add_hyperlink_run(p, article["keyword2"], article["url2"])
        elif MARKER_RE.fullmatch(part or ""):
            continue              # 未對應嘅 marker 唔可以印落成品
        elif part:
            _set_run_font(p.add_run(part))
    return p


def _add_heading(doc, level, title):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"[H{level}: {title}]")
    _set_run_font(run, size=17 if level == 1 else 13, bold=True, color="000000")
    return p


def _init_doc():
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Arial"
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return doc


def _write_article(doc, article, content, with_header=True):
    if with_header:
        p = doc.add_paragraph()
        _set_run_font(p.add_run(f"#{article['number']}"), bold=True)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _set_run_font(p.add_run("Keyword："), bold=True)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _set_run_font(p.add_run(f"● {article['keyword1']}"))
        if article["keyword2"]:
            p = doc.add_paragraph()
            _set_run_font(p.add_run(f"● {article['keyword2']}"))
        doc.add_paragraph()

    _add_heading(doc, 1, content["h1"])

    for section in content["sections"]:
        if section.get("h2"):
            _add_heading(doc, 2, section["h2"])
        body = section.get("body", "")
        if body:
            _add_body_paragraph(doc, body, article)


def build_single_docx(article, content, output_path):
    """一篇一個檔,冇 #N / Keyword 頭段 ——
    可以直接餵去 dz-linkbuild 嘅 validate.py 做最終覆核。"""
    doc = _init_doc()
    _write_article(doc, article, content, with_header=False)
    doc.save(output_path)
    return output_path


def build_docx_file(articles_with_content, output_path):
    """合併檔:每篇之間分頁,保留 #N / Keyword 頭段(DZ 交付慣例)。

    注意:validate.py 係為單篇檔設計,#N 同 ● keyword 行會被當成正文,
    令關鍵字次數統計多咗。要跑 validator 請用 build_single_docx()。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _init_doc()
    for idx, (article, content) in enumerate(articles_with_content):
        _write_article(doc, article, content, with_header=True)
        if idx < len(articles_with_content) - 1:
            p = doc.add_paragraph()
            run = p.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._element.append(br)

    doc.save(output_path)
    return output_path


# ================================================================
# Google Docs API
# ================================================================
def get_google_services(credentials_path):
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    creds = service_account.Credentials.from_service_account_file(
        credentials_path, scopes=SCOPES
    )
    return build("docs", "v1", credentials=creds), build("drive", "v3", credentials=creds)


def get_google_services_from_info(credentials_info):
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    info = dict(credentials_info)
    if "private_key" in info:
        info["private_key"] = info["private_key"].replace("\\n", "\n")
    creds = service_account.Credentials.from_service_account_info(info, scopes=SCOPES)
    return build("docs", "v1", credentials=creds), build("drive", "v3", credentials=creds)


def create_formatted_doc(
    docs_service, drive_service, builder, title,
    folder_id=None, share_email=None
):
    doc = docs_service.documents().create(body={"title": title}).execute()
    doc_id = doc["documentId"]
    print(f"  📄 Created doc: {doc_id}")

    requests = builder.get_gdocs_requests()
    chunk_size = 80
    for i in range(0, len(requests), chunk_size):
        docs_service.documents().batchUpdate(
            documentId=doc_id, body={"requests": requests[i:i + chunk_size]},
        ).execute()

    if folder_id:
        file_meta = drive_service.files().get(fileId=doc_id, fields="parents").execute()
        prev_parents = ",".join(file_meta.get("parents", []))
        drive_service.files().update(
            fileId=doc_id, addParents=folder_id,
            removeParents=prev_parents, fields="id, parents",
        ).execute()
        print(f"  📁 Moved to folder: {folder_id}")

    if share_email:
        drive_service.permissions().create(
            fileId=doc_id,
            body={"type": "user", "role": "writer", "emailAddress": share_email},
            sendNotificationEmail=False,
        ).execute()
        print(f"  👤 Shared with: {share_email}")

    return doc_id, f"https://docs.google.com/document/d/{doc_id}/edit"


# ================================================================
# Main Pipeline
# ================================================================
def main():
    parser = argparse.ArgumentParser(
        description="Digital Zoo Linkbuilding Article Generator"
    )
    parser.add_argument("--excel", required=True, help="Excel 檔案路徑")
    parser.add_argument("--batch", required=True, type=int, help="Batch 編號 (1-4)")
    parser.add_argument("--creds", default=None, help="Google Service Account JSON")
    parser.add_argument("--folder", default=None, help="Google Drive 資料夾 ID")
    parser.add_argument("--share", default=None, help="分享給指定 email")
    parser.add_argument("--start", type=int, default=None, help="從第 N 篇開始")
    parser.add_argument("--end", type=int, default=None, help="到第 N 篇結束")
    parser.add_argument("--sheet", default=None, help="Excel sheet 名稱(預設第一個 sheet)")
    parser.add_argument("--docx", default=None, help="輸出合併 .docx 到指定路徑")
    parser.add_argument("--single-docx-dir", default=None,
                        help="每篇輸出一個 .docx 到指定資料夾(可直接跑 validate.py)")
    parser.add_argument("--dry-run", action="store_true", help="只生成文字唔建檔")
    args = parser.parse_args()

    api_key = OPENROUTER_API_KEY
    if not api_key:
        print("✗ 請設定環境變數 OPENROUTER_API_KEY")
        sys.exit(1)

    print(f"\n📊 讀取 Excel: {args.excel} (Batch {args.batch})")
    sheet_warnings = []
    articles = parse_excel(args.excel, args.batch, args.sheet, report=sheet_warnings)

    for w in sheet_warnings:
        print(f"  ⚠ {w}")

    if not articles:
        print("✗ 找不到指定 Batch 的文章")
        sys.exit(1)

    if args.start:
        articles = [a for a in articles if a["number"] >= args.start]
    if args.end:
        articles = [a for a in articles if a["number"] <= args.end]

    print(f"  找到 {len(articles)} 篇文章 (#{articles[0]['number']}-#{articles[-1]['number']})")
    for a in articles:
        for w in a.get("warnings", []):
            print(f"  ⚠ #{a['number']}: {w}")

    print(f"\n🤖 開始生成文章 (Model: {MODEL})")
    builder = DocBuilder()
    done, failed, noncompliant = [], [], []

    for article in tqdm(articles, desc="生成進度"):
        num = article["number"]
        kw1, kw2 = article["keyword1"], article["keyword2"]
        tqdm.write(f"  #{num}: {kw1}" + (f" + {kw2}" if kw2 else ""))

        content = generate_article_content(article, api_key, MODEL)
        if content and not content.get("_compliant", True):
            noncompliant.append((article, content))
            tqdm.write(f"  ⚠ #{num} 未合規({len(content.get('_fails', []))} 項),已隔開")
            content = None
        if content:
            builder.build_article(article, content)
            done.append((article, content))
            tqdm.write(
                f"  ✓ #{num} 完成 ({content['_word_count']} "
                f"{'words' if content['_lang'] == 'en' else '字'}, 合規通過)"
            )
            for w in content.get("_warnings", []):
                tqdm.write(f"     ⚠ {w}")
        else:
            failed.append(num)
            tqdm.write(f"  ✗ #{num} 失敗")

        time.sleep(DELAY_BETWEEN_CALLS)

    print(f"\n📝 生成完成: {len(done)}/{len(articles)} 篇")
    if noncompliant:
        print(f"  ⚠ 未合規(需人手處理): {[a['number'] for a, _ in noncompliant]}")
        for a, c in noncompliant:
            for f in c.get("_fails", []):
                print(f"      #{a['number']} • {f}")
    if failed:
        print(f"  ⚠ 完全失敗: {failed}")

    placeholders = [
        a["number"] for a, c in done
        if PLACEHOLDER_URL in (a["url1"], a["url2"])
    ]
    if placeholders:
        print(f"  ⚠ 仲用緊 placeholder URL 嘅文章: {placeholders}")

    if args.single_docx_dir:
        out_dir = Path(args.single_docx_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        for article, content in done:
            path = out_dir / f"{article['number']:03d}_{article['keyword1'][:20]}.docx"
            build_single_docx(article, content, str(path))
        print(f"\n💾 單篇 .docx 已輸出到: {out_dir}")

    if args.docx:
        build_docx_file(done, args.docx)
        print(f"\n💾 合併 .docx: {args.docx}")

    if args.dry_run:
        out_path = f"batch_{args.batch}_output.txt"
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(builder.text)
        print(f"\n💾 Dry run 輸出: {out_path}")
        return

    if not args.creds:
        if not (args.docx or args.single_docx_dir):
            print("\n⚠ 冇指定 --creds / --docx / --single-docx-dir,冇建立任何檔案")
        return

    if not Path(args.creds).exists():
        print(f"✗ 找不到 credentials 檔案: {args.creds}")
        sys.exit(1)

    print("\n📄 建立 Google Doc...")
    docs_svc, drive_svc = get_google_services(args.creds)

    title = f"Combined_2026_May_Internal_Batch_{args.batch}"
    if args.start or args.end:
        title += f"_#{args.start or articles[0]['number']}-{args.end or articles[-1]['number']}"

    doc_id, doc_url = create_formatted_doc(
        docs_svc, drive_svc, builder, title,
        folder_id=args.folder, share_email=args.share,
    )

    print("\n✅ 完成！")
    print(f"   文件名: {title}")
    print(f"   連結: {doc_url}")


if __name__ == "__main__":
    main()
